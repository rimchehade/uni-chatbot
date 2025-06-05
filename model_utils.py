import os
import json
import torch
import librosa
import soundfile as sf
from fpdf import FPDF
from docx import Document
from duckduckgo_search import DDGS
from transformers import pipeline, AutoTokenizer, AutoModelForCausalLM, BitsAndBytesConfig
from langchain.vectorstores import FAISS
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.memory import ConversationBufferMemory
from langchain.chains import ConversationalRetrievalChain
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain
from langchain_huggingface import HuggingFacePipeline
import ffmpeg


# ========== Initialize Whisper ==========
asr_pipeline = pipeline(
    "automatic-speech-recognition",
    model="openai/whisper-small",
    tokenizer="openai/whisper-small",
    generate_kwargs={"language": "en"},
    return_timestamps=False
)


def transcribe_audio(audio_path):
    audio, sr = librosa.load(audio_path, sr=16000)
    sf.write("converted.wav", audio, 16000)
    result = asr_pipeline("converted.wav")
    return result["text"]
    

# ========== Initialize Embeddings and Vectorstore ==========
embedding_model = HuggingFaceEmbeddings(
    model_name="BAAI/bge-large-en-v1.5",
    encode_kwargs={"normalize_embeddings": True}
)

vectorstore = FAISS.load_local(
    "lau_vectorstore",
    embeddings=embedding_model,
    allow_dangerous_deserialization=True
)

retriever = vectorstore.as_retriever(search_kwargs={"k": 3})


# ========== Initialize LLM ==========
def initialize_llm():
    model_name = "meta-llama/Llama-2-13b-chat-hf"
    hf_token = os.environ.get("HF_TOKEN")
    device = "cuda" if torch.cuda.is_available() else "cpu"

    tokenizer = AutoTokenizer.from_pretrained(model_name, token=hf_token)

    bnb_config = BitsAndBytesConfig(load_in_4bit=True,
                                    bnb_4bit_compute_dtype=torch.float16)

    model = AutoModelForCausalLM.from_pretrained(
        model_name,
        token=hf_token,
        quantization_config=bnb_config
    ).to(device)

    pipe = pipeline(
        "text-generation",
        model=model,
        tokenizer=tokenizer,
        pad_token_id=tokenizer.eos_token_id,
        temperature=0.7,
        top_p=0.9,
        max_new_tokens=512,
        do_sample=True
    )

    return HuggingFacePipeline(pipeline=pipe)


llm = initialize_llm()

# ========== Chains ==========
follow_up_prompt = PromptTemplate(
    input_variables=["question", "answer"],
    template="""
You are a helpful assistant. Given the user's question and your answer, suggest a relevant follow-up question.

Question: {question}
Answer: {answer}

Follow-up question:
"""
)

follow_up_chain = LLMChain(llm=llm, prompt=follow_up_prompt)

memory = ConversationBufferMemory(memory_key="chat_history", return_messages=True, output_key="answer")

qa_chain = ConversationalRetrievalChain.from_llm(
    llm=llm,
    retriever=retriever,
    memory=memory,
    return_source_documents=True
)


# ========== Core Chatbot Function ==========
def answer_question(question):
    result = qa_chain.invoke({"question": question})
    answer = result["answer"]
    sources = result.get("source_documents", [])

    context = "\n\n".join(doc.page_content for doc in sources)
    decision_prompt = f"""
Only return "1" if the context directly answers the question.
If not, return "0".

Context: {context}
Question: {question}
Answer:
"""

    decision = llm.predict(decision_prompt).strip()
    follow_up = follow_up_chain.run({"question": question, "answer": answer})

    return {
        "question": question,
        "answer": answer,
        "follow_up": follow_up,
        "decision": decision
    }


# ========== Public Interfaces ==========
def handle_text_input(question):
    return answer_question(question)


def handle_audio_upload(audio_path):
    question = transcribe_audio(audio_path)
    return answer_question(question)


def save_answer_to_pdf(question, answer, filename="answer.pdf"):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, f"Question:\n{question}\n\nAnswer:\n{answer}")
    pdf.output(filename)


def save_conversation_to_docx(conversation, filename="chat.docx"):
    doc = Document()
    doc.add_heading("Conversation Log", 0)
    for i, (q, a) in enumerate(conversation, 1):
        doc.add_heading(f"Q{i}:", level=1)
        doc.add_paragraph(q)
        doc.add_heading(f"A{i}:", level=2)
        doc.add_paragraph(a)
    doc.save(filename)
