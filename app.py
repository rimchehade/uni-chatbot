import streamlit as st
import requests
from fpdf import FPDF
from docx import Document
import base64
import io

st.set_page_config(page_title="Uni Chat", layout="centered")
st.title("🎓 Uni Chat – Lebanese Universities Chatbot")

if "history" not in st.session_state:
    st.session_state.history = []

if "bookmarks" not in st.session_state:
    st.session_state.bookmarks = []

# --- Audio Recording ---
st.subheader("Ask via Voice or Text")
audio_data = st.file_uploader("Record or upload audio", type=["wav", "mp3", "m4a"])

user_input = st.text_input("Or type your question:")

# --- Send Query ---
if st.button("Send"):
    query = None

    if audio_data:
        # Assume your backend accepts audio as base64
        b64_audio = base64.b64encode(audio_data.read()).decode("utf-8")
        response = requests.post("https://your-backend-url.com/audio", json={"audio": b64_audio})
        query = "[Voice message]"
    elif user_input:
        response = requests.post("https://your-backend-url.com/text", json={"query": user_input})
        query = user_input
    else:
        st.warning("Please provide text or audio input.")
        response = None

    if response and response.ok:
        answer = response.json().get("answer", "No response.")
        st.session_state.history.append({"user": query, "bot": answer})
    elif response:
        st.error("Something went wrong. Please try again.")

# --- Display Chat ---
for idx, chat in enumerate(st.session_state.history):
    st.markdown(f"**You:** {chat['user']}")
    st.markdown(f"**Uni Chat:** {chat['bot']}")
    col1, col2 = st.columns([1, 1])
    with col1:
        if st.button("🔖 Bookmark", key=f"bookmark_{idx}"):
            st.session_state.bookmarks.append(chat)
    with col2:
        st.download_button("⬇️ Export",
                           data=chat['bot'],
                           file_name=f"uni_chat_{idx}.txt",
                           mime="text/plain",
                           key=f"export_{idx}")

# --- Export Full Chat ---
st.markdown("### 📄 Export Full Chat")

col_pdf, col_docx = st.columns(2)

with col_pdf:
    if st.button("Export as PDF"):
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        for chat in st.session_state.history:
            pdf.multi_cell(0, 10, f"You: {chat['user']}\nUni Chat: {chat['bot']}\n")
        buffer = io.BytesIO()
        pdf.output(buffer)
        st.download_button("Download PDF", data=buffer.getvalue(), file_name="chat_history.pdf", mime="application/pdf")

with col_docx:
    if st.button("Export as Word"):
        doc = Document()
        for chat in st.session_state.history:
            doc.add_paragraph(f"You: {chat['user']}")
            doc.add_paragraph(f"Uni Chat: {chat['bot']}")
        buffer = io.BytesIO()
        doc.save(buffer)
        st.download_button("Download Word Doc", data=buffer.getvalue(), file_name="chat_history.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
